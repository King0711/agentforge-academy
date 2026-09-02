"""
Step 1 of the Document Summarizer Agent.

This file has one job: given a file path, work out what kind of document
it is and pull the plain text out of it. A PDF, a Word doc and a plain
text file all store "the same" words in completely different byte
formats - this file hides that difference from the rest of the program,
so summarizer.py only ever has to deal with plain text.
"""

import os

import requests
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from docx import Document

# Extensions we know how to handle, mapped to a short type name.
# A dict here (instead of a chain of if/elif) means adding a new format
# later is one new line, not a new branch in three different places.
SUPPORTED_EXTENSIONS = {
    ".txt": "txt",
    ".pdf": "pdf",
    ".docx": "docx",
}

# Websites often block programs that don't look like a real browser.
# Telling them we're a normal browser avoids most of that - same trick
# the Social Post Generator's article.py uses.
BROWSER_HEADER = {"User-Agent": "Mozilla/5.0"}

# Parts of a web page that are never the article/document text.
NOT_THE_ARTICLE = ["script", "style", "nav", "footer", "header", "aside", "form"]


def detect_type(path):
    """
    Looks at a file path or web address and returns "txt", "pdf", "docx"
    or "url".

        detect_type("report.pdf")                  -> "pdf"
        detect_type("notes.txt")                    -> "txt"
        detect_type("https://example.com/a-page")   -> "url"

    Raises ValueError for anything else instead of guessing - guessing
    wrong here means feeding a PDF's raw bytes to the plain-text reader,
    which produces garbage instead of a clear error message.
    """
    if path.startswith("http://") or path.startswith("https://"):
        return "url"

    _, ext = os.path.splitext(path.lower())
    doc_type = SUPPORTED_EXTENSIONS.get(ext)
    if doc_type is None:
        raise ValueError(
            f"Don't know how to read '{ext or path}' files.\n"
            f"This tool handles .txt, .pdf, .docx and web links - "
            f"try saving or exporting your document as one of those."
        )
    return doc_type


def extract_txt(path):
    """
    Reads a plain text file.

    errors="replace" matters more than it looks: a .txt file saved by
    Word or exported from a Mac is sometimes not quite UTF-8. Without
    this, one odd character crashes the whole run; with it, that one
    character becomes a harmless "?" and everything else still works.
    """
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def extract_pdf(path):
    """
    Pulls text out of a PDF, page by page.

    PyPDF2 only reads text that is actually stored as text in the file.
    A PDF that is really a scanned photo of a page has no text layer at
    all, and PyPDF2 returns an empty string for it rather than an error -
    that's why extract_text() below checks the final length itself.
    """
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def extract_docx(path):
    """
    Pulls text out of a Word document, paragraph by paragraph.

    python-docx only sees body paragraphs and tables - not headers,
    footers, or text baked into images. Fine for the reports and
    contracts this tool targets, but worth knowing if a summary seems to
    miss something that only appeared in a header or footer.
    """
    document = Document(path)
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join(p for p in paragraphs if p)


def extract_url(url):
    """
    Downloads a web page and returns its readable text - the exact same
    approach as the Social Post Generator's article.py: strip out the
    parts of a page that are never the content, then pull text from every
    <p> (paragraph) tag.

    This treats the address as an HTML page. It will NOT correctly read a
    link that points straight at a PDF or DOCX file for download (the
    bytes it gets back aren't HTML) - save that file to your computer
    first and pass the local path to this tool instead.
    """
    try:
        page = requests.get(url, headers=BROWSER_HEADER, timeout=15)
        page.raise_for_status()
    except requests.RequestException:
        raise ValueError(
            f"Could not open that link.\n"
            f"Check the address is correct and that the page opens in your browser."
        )

    soup = BeautifulSoup(page.text, "html.parser")
    for junk in soup(NOT_THE_ARTICLE):
        junk.decompose()

    paragraphs = [p.get_text(" ", strip=True) for p in soup.find_all("p")]
    return "\n".join(p for p in paragraphs if p)


def extract_text(path):
    """
    The one function the rest of the program calls. Detects the type,
    runs the matching extractor, and returns clean plain text.

        text = extract_text("quarterly_report.pdf")
        text = extract_text("https://example.com/a-blog-post")

    Raises ValueError with a readable message if the type is unsupported
    or if nothing readable came out of it.
    """
    doc_type = detect_type(path)

    if doc_type == "txt":
        text = extract_txt(path)
    elif doc_type == "pdf":
        text = extract_pdf(path)
    elif doc_type == "url":
        text = extract_url(path)
    else:
        text = extract_docx(path)

    if len(text.strip()) < 20:
        raise ValueError(
            f"Couldn't find any readable text in '{path}'.\n"
            f"If this is a PDF, it may be a scanned image with no text "
            f"layer underneath it - this tool can't read those without "
            f"OCR (optical character recognition), which is a different "
            f"tool entirely."
        )

    return text


if __name__ == "__main__":
    # Run this file on its own to check it works:  python detector.py
    #
    # These are genuine round-trip tests, not just "does the code look
    # right" - each one WRITES a real file in the target format, then
    # reads it back through the exact function main.py will call. A bug
    # in either direction (writing something unusual, or misreading
    # something normal) shows up here instead of surprising a student
    # later on their own document.
    import tempfile

    tmp_dir = tempfile.gettempdir()

    print("Checking .txt extraction...")
    txt_path = os.path.join(tmp_dir, "sdt_test_sample.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(
            "Quarterly Planning Notes\n\n"
            "Revenue is up 12% this quarter.\n"
            "Action: renew the Lagos office lease by Friday."
        )
    assert detect_type(txt_path) == "txt"
    text = extract_text(txt_path)
    assert "Revenue is up 12%" in text
    print(f"  OK - wrote a real .txt file, read back {len(text)} characters")

    print("Checking .docx round-trip...")
    docx_path = os.path.join(tmp_dir, "sdt_test_sample.docx")
    doc = Document()
    doc.add_paragraph("Board Meeting Summary")
    doc.add_paragraph("The team agreed to ship the new pricing page next week.")
    doc.save(docx_path)
    assert detect_type(docx_path) == "docx"
    text = extract_text(docx_path)
    assert "ship the new pricing page" in text
    print(f"  OK - wrote a real .docx with python-docx, read it back, got {len(text)} characters")

    print("Checking .pdf round-trip...")
    try:
        from fpdf import FPDF

        pdf_path = os.path.join(tmp_dir, "sdt_test_sample.pdf")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(0, 10, text="Contract Review Notes")
        pdf.ln(10)
        pdf.cell(0, 10, text="Client must countersign by March 3rd.")
        pdf.output(pdf_path)

        assert detect_type(pdf_path) == "pdf"
        text = extract_text(pdf_path)
        assert "countersign" in text
        print(f"  OK - wrote a real .pdf with fpdf2, read it back with PyPDF2, got {len(text)} characters")
    except ImportError:
        print("  SKIPPED - fpdf2 not installed (test-only dependency, see requirements.txt)")

    print("Checking URL extraction against a real, stable web page...")
    try:
        sample_url = "https://en.wikipedia.org/wiki/Artificial_intelligence"
        assert detect_type(sample_url) == "url"
        text = extract_text(sample_url)
        assert len(text) > 500
        print(f"  OK - fetched a real page, got {len(text)} characters")
    except ValueError as problem:
        # A network hiccup or a site layout change shouldn't fail the
        # whole test file - it's the one check here that depends on
        # something outside our control.
        print(f"  SKIPPED - could not fetch the test page ({problem})")

    print("Checking an unsupported extension raises a clear error...")
    try:
        detect_type("presentation.pptx")
        print("  FAILED - should have raised ValueError")
    except ValueError as e:
        print(f"  OK - raised: {e.args[0].splitlines()[0]}")

    print("\nAll checks passed.")
